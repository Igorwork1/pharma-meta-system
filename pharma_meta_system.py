import streamlit as st
import pandas as pd
import psycopg
import logging
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
import io
import uuid
import re
from docx import Document
from docx.shared import Inches
import pdfkit
from tempfile import NamedTemporaryFile
import base64
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx import Document

# Настройка логирования
logging.basicConfig(
    filename='pharma_metadata.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
def log_action(action, details=None, username=None):
    log_msg = f"{action} {'by ' + username if username else ''}"
    if details:
        log_msg += f" - Details: {details}"
    log_msg = f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - INFO - {log_msg}\n"
    log_file = 'edit_access_denied.log' if action == "Access denied to edit data" else 'pharma_metadata.log'
    try:
        with open(log_file, 'a') as f:
            f.write(log_msg)
    except Exception as e:
        print(f"Error writing to log {log_file}: {e}")

def get_logs(log_type='main'):
    log_file = 'pharma_metadata.log' if log_type == 'main' else 'edit_access_denied.log'
    try:
        with open(log_file, 'r') as f:
            return f.readlines()
    except FileNotFoundError:
        return [f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - INFO - Log file {log_file} created\n"]
    
def clear_logs_daily():
    import os
    import time
    log_files = ['pharma_metadata.log', 'edit_access_denied.log']
    twenty_four_hours = 24 * 60 * 60  # 24 hours in seconds
    current_time = time.time()
    for log_file in log_files:
        try:
            if os.path.exists(log_file):
                mtime = os.path.getmtime(log_file)
                if current_time - mtime > twenty_four_hours:
                    with open(log_file, 'w') as f:
                        f.write(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - INFO - Log cleared\n")
        except Exception as e:
            print(f"Error clearing log {log_file}: {e}")
            
# Настройка подключения к PostgreSQL
def get_db_connection():
    try:
        conn = psycopg.connect(
            dbname="meta_base",
            user="postgres",
            password="1234",
            host="localhost",
            port="5432",
            sslmode="require"  # Для облачной базы, если требуется
        )
        return conn
    except psycopg.Error as e:
        st.error(f"Ошибка подключения к базе данных: {e}")
        return None
        
def init_db():
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    # Создание таблицы companies
    c.execute('''CREATE TABLE IF NOT EXISTS companies (
        id SERIAL PRIMARY KEY,
        gln VARCHAR(20),
        name_short VARCHAR(50),
        name_full VARCHAR(100),
        gcp_compliant BOOLEAN,
        registration_country VARCHAR(50),
        address VARCHAR(200),
        type VARCHAR(50)
    )''')
    # Создание таблицы medicines
    c.execute('''CREATE TABLE IF NOT EXISTS medicines (
        id SERIAL PRIMARY KEY,
        owned_by INTEGER,
        name VARCHAR(50),
        gtin VARCHAR(20),
        sku VARCHAR(20),
        market VARCHAR(20),
        shared BOOLEAN DEFAULT FALSE,
        batch_number VARCHAR(50),
        expiration_date DATE,
        dosage_form VARCHAR(50),
        active_ingredient VARCHAR(100),
        package_size VARCHAR(50),
        created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (owned_by) REFERENCES companies(id) ON DELETE SET NULL
    )''')
    # Создание таблицы locations
    c.execute('''CREATE TABLE IF NOT EXISTS locations (
        id SERIAL PRIMARY KEY,
        owned_by INTEGER,
        gln VARCHAR(20),
        country VARCHAR(50),
        address VARCHAR(200),
        role VARCHAR(50),
        name_short VARCHAR(50),
        name_full VARCHAR(100),
        created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (owned_by) REFERENCES companies(id) ON DELETE SET NULL
    )''')
    # Создание таблицы operations
    c.execute('''CREATE TABLE IF NOT EXISTS operations (
        id SERIAL PRIMARY KEY,
        medicine_id INTEGER,
        location_id INTEGER,
        operation_type VARCHAR(50),
        operation_date TIMESTAMP,
        quantity INTEGER,
        created_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (medicine_id) REFERENCES medicines(id) ON DELETE SET NULL,
        FOREIGN KEY (location_id) REFERENCES locations(id) ON DELETE SET NULL
    )''')
    # Создание таблицы users
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        id SERIAL PRIMARY KEY,
        login VARCHAR(20) UNIQUE,
        password VARCHAR(50),
        role VARCHAR(20),
        first_name VARCHAR(50),
        last_name VARCHAR(50),
        email VARCHAR(100)
    )''')
    # Добавление столбца medicine_id, если он отсутствует
    c.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'operations' AND column_name = 'medicine_id'")
    if not c.fetchone():
        c.execute("ALTER TABLE operations ADD COLUMN medicine_id INTEGER REFERENCES medicines(id) ON DELETE SET NULL")

    # Добавление столбца location_id, если он отсутствует
    c.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'operations' AND column_name = 'location_id'")
    if not c.fetchone():
        c.execute("ALTER TABLE operations ADD COLUMN location_id INTEGER REFERENCES locations(id) ON DELETE SET NULL")

    # Удаление ограничений UNIQUE для gtin и sku, если они существуют
    c.execute("SELECT constraint_name FROM information_schema.table_constraints WHERE table_name = 'medicines' AND constraint_type = 'UNIQUE' AND constraint_name LIKE '%gtin%'")
    gtin_constraint = c.fetchone()
    if gtin_constraint:
        c.execute(f"ALTER TABLE medicines DROP CONSTRAINT {gtin_constraint[0]}")

    c.execute("SELECT constraint_name FROM information_schema.table_constraints WHERE table_name = 'medicines' AND constraint_type = 'UNIQUE' AND constraint_name LIKE '%sku%'")
    sku_constraint = c.fetchone()
    if sku_constraint:
        c.execute(f"ALTER TABLE medicines DROP CONSTRAINT {sku_constraint[0]}")
        
    # Добавляем столбец atc_code, если его нет
    c.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'medicines' AND column_name = 'atc_code'")
    if not c.fetchone():
        c.execute("ALTER TABLE medicines ADD COLUMN atc_code VARCHAR(20)")

    conn.commit()
    conn.close()

# Функции получения данных
def get_medications():
    conn = get_db_connection()
    if conn is None:
        return pd.DataFrame()
    df = pd.read_sql_query("SELECT * FROM medicines", conn)
    conn.close()
    return df

def get_companies():
    conn = get_db_connection()
    if conn is None:
        return pd.DataFrame()
    df = pd.read_sql_query("SELECT * FROM companies", conn)
    conn.close()
    return df

def get_locations():
    conn = get_db_connection()
    if conn is None:
        return pd.DataFrame()
    df = pd.read_sql_query("SELECT * FROM locations", conn)
    conn.close()
    return df

def get_operations():
    conn = get_db_connection()
    if conn is None:
        return pd.DataFrame()
    df = pd.read_sql_query("SELECT * FROM operations", conn)
    conn.close()
    return df

# Функции добавления
def add_medication(name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''INSERT INTO medicines 
                     (name, gtin, sku, market, shared, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code, created_date) 
                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)''',
                  (name, gtin, sku, market, False, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code,
                   datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
        log_action("Added medication", f"ID: {c.lastrowid}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка добавления Препарата: {e}")
    finally:
        conn.close()

def add_company(gln, name_short, name_full, gcp_compliant, registration_country, address, type, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''INSERT INTO companies 
                     (gln, name_short, name_full, gcp_compliant, registration_country, address, type) 
                     VALUES (%s, %s, %s, %s, %s, %s, %s)''',
                  (gln, name_short, name_full, gcp_compliant, registration_country, address, type))
        conn.commit()
        log_action("Added company", f"ID: {c.lastrowid}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка добавления компании: {e}")
    finally:
        conn.close()

def add_location(gln, country, address, role, name_short, name_full, owned_by, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''INSERT INTO locations 
                     (gln, country, address, role, name_short, name_full, owned_by, created_date) 
                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s)''',
                  (gln, country, address, role, name_short, name_full, owned_by,
                   datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
        log_action("Added location", f"ID: {c.lastrowid}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка добавления локации: {e}")
    finally:
        conn.close()

def add_operation(medicine_id, location_id, operation_type, operation_date, quantity, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''INSERT INTO operations 
                     (medicine_id, location_id, operation_type, operation_date, quantity, created_date) 
                     VALUES (%s, %s, %s, %s, %s, %s)''',
                  (medicine_id, location_id, operation_type, operation_date, quantity,
                   datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
        log_action("Added operation", f"ID: {c.lastrowid}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка добавления операции: {e}")
    finally:
        conn.close()

# Функции редактирования
def edit_medication(med_id, name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''UPDATE medicines 
                     SET name=%s, gtin=%s, sku=%s, market=%s, batch_number=%s, expiration_date=%s, 
                         dosage_form=%s, active_ingredient=%s, package_size=%s, owned_by=%s, atc_code=%s 
                     WHERE id=%s''',
                  (name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code, med_id))
        conn.commit()
        log_action("Edited medication", f"ID: {med_id}, Changed fields: {', '.join([f'{k}={v}' for k, v in {'name': name, 'gtin': gtin, 'sku': sku, 'market': market, 'batch_number': batch_number, 'expiration_date': str(expiration_date), 'dosage_form': dosage_form, 'active_ingredient': active_ingredient, 'package_size': package_size, 'owned_by': owned_by, 'atc_code': atc_code}.items() if v])}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка редактирования Препарата: {e}")
    finally:
        conn.close()

def edit_company(company_id, gln, name_short, name_full, gcp_compliant, registration_country, address, type, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''UPDATE companies 
                     SET gln=%s, name_short=%s, name_full=%s, gcp_compliant=%s, registration_country=%s, address=%s, type=%s 
                     WHERE id=%s''',
                  (gln, name_short, name_full, gcp_compliant, registration_country, address, type, company_id))
        conn.commit()
        log_action("Edited company", f"ID: {company_id}, Changed fields: {', '.join([f'{k}={v}' for k, v in {'gln': gln, 'name_short': name_short, 'name_full': name_full, 'gcp_compliant': str(gcp_compliant), 'registration_country': registration_country, 'address': address, 'type': type}.items() if v])}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка редактирования компании: {e}")
    finally:
        conn.close()

def edit_location(location_id, gln, country, address, role, name_short, name_full, owned_by, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''UPDATE locations 
                     SET gln=%s, country=%s, address=%s, role=%s, name_short=%s, name_full=%s, owned_by=%s 
                     WHERE id=%s''',
                  (gln, country, address, role, name_short, name_full, owned_by, location_id))
        conn.commit()
        log_action("Edited location", f"ID: {location_id}, Changed fields: {', '.join([f'{k}={v}' for k, v in {'gln': gln, 'country': country, 'address': address, 'role': role, 'name_short': name_short, 'name_full': name_full, 'owned_by': owned_by}.items() if v])}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка редактирования локации: {e}")
    finally:
        conn.close()

def edit_operation(operation_id, medicine_id, location_id, operation_type, operation_date, quantity, username):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute('''UPDATE operations 
                     SET medicine_id=%s, location_id=%s, operation_type=%s, operation_date=%s, quantity=%s 
                     WHERE id=%s''',
                  (medicine_id, location_id, operation_type, operation_date, quantity, operation_id))
        conn.commit()
        log_action("Edited operation", f"ID: {operation_id}, Changed fields: {', '.join([f'{k}={v}' for k, v in {'medicine_id': medicine_id, 'location_id': location_id, 'operation_type': operation_type, 'operation_date': str(operation_date), 'quantity': quantity}.items() if v])}", username)
    except psycopg2.Error as e:
        st.error(f"Ошибка редактирования операции: {e}")
    finally:
        conn.close()

# Функции удаления с проверкой зависимостей
def delete_medication(med_id):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute("SELECT COUNT(*) FROM operations WHERE medicine_id = %s", (med_id,))
        if c.fetchone()[0] > 0:
            st.error("Нельзя удалить Препарат, так как он связан с операциями")
            return
        c.execute("DELETE FROM medicines WHERE id=%s", (med_id,))
        conn.commit()
        log_action("Deleted medication", f"ID: {med_id}")
    except psycopg2.Error as e:
        st.error(f"Ошибка удаления Препарата: {e}")
    finally:
        conn.close()

def delete_company(company_id):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute("SELECT COUNT(*) FROM medicines WHERE owned_by = %s", (company_id,))
        med_count = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM locations WHERE owned_by = %s", (company_id,))
        loc_count = c.fetchone()[0]
        if med_count > 0 or loc_count > 0:
            st.error("Нельзя удалить компанию, так как у неё есть связанные Препараты или локации")
            return
        c.execute("DELETE FROM companies WHERE id=%s", (company_id,))
        conn.commit()
        log_action("Deleted company", f"ID: {company_id}")
    except psycopg2.Error as e:
        st.error(f"Ошибка удаления компании: {e}")
    finally:
        conn.close()

def delete_location(location_id):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute("SELECT COUNT(*) FROM operations WHERE location_id = %s", (location_id,))
        if c.fetchone()[0] > 0:
            st.error("Нельзя удалить локацию, так как она связана с операциями")
            return
        c.execute("DELETE FROM locations WHERE id=%s", (location_id,))
        conn.commit()
        log_action("Deleted location", f"ID: {location_id}")
    except psycopg2.Error as e:
        st.error(f"Ошибка удаления локации: {e}")
    finally:
        conn.close()

def delete_operation(operation_id):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        c.execute("DELETE FROM operations WHERE id=%s", (operation_id,))
        conn.commit()
        log_action("Deleted operation", f"ID: {operation_id}")
    except psycopg2.Error as e:
        st.error(f"Ошибка удаления операции: {e}")
    finally:
        conn.close()

# Импорт/экспорт данных
def import_data(file):
    conn = get_db_connection()
    if conn is None:
        return
    c = conn.cursor()
    try:
        if file.type in ['text/csv', 'application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            df = pd.read_csv(file) if file.type == 'text/csv' else pd.read_excel(file)
            table = df.columns[0].split('_')[0]
            if table == 'medicines':
                for _, row in df.iterrows():
                    c.execute('''SELECT id FROM medicines 
                                 WHERE name = %s AND gtin = %s AND sku = %s AND market = %s 
                                 AND batch_number = %s AND expiration_date = %s AND dosage_form = %s 
                                 AND active_ingredient = %s AND package_size = %s AND owned_by = %s''AND atc_code = %s''',
                              (row['name'], row['gtin'], row['sku'], row['market'], row['batch_number'],
                               row['expiration_date'], row['dosage_form'], row['active_ingredient'], row['package_size'], row['owned_by'], row.get('atc_code', None)))
                    if c.fetchone():
                        continue
                    c.execute('''INSERT INTO medicines 
                                 (name, gtin, sku, market, shared, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, created_date) 
                                 VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)''',
                              (row['name'], row['gtin'], row['sku'], row['market'], row['shared'], row['batch_number'],
                               row['expiration_date'], row['dosage_form'], row['active_ingredient'], row['package_size'], row['owned_by'], datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            elif table == 'companies':
                for _, row in df.iterrows():
                    c.execute('''SELECT id FROM companies 
                                 WHERE gln = %s AND name_short = %s AND name_full = %s 
                                 AND gcp_compliant = %s AND registration_country = %s 
                                 AND address = %s AND type = %s''',
                              (row['gln'], row['name_short'], row['name_full'], row['gcp_compliant'],
                               row['registration_country'], row['address'], row['type']))
                    if c.fetchone():
                        continue
                    c.execute('''INSERT INTO companies 
                                 (gln, name_short, name_full, gcp_compliant, registration_country, address, type) 
                                 VALUES (%s, %s, %s, %s, %s, %s, %s)''',
                              (row['gln'], row['name_short'], row['name_full'], row['gcp_compliant'],
                               row['registration_country'], row['address'], row['type']))
            elif table == 'locations':
                for _, row in df.iterrows():
                    c.execute('''INSERT INTO locations 
                                 (gln, country, address, role, name_short, name_full, owned_by, created_date) 
                                 VALUES (%s, %s, %s, %s, %s, %s, %s, %s)''',
                              (row['gln'], row['country'], row['address'], row['role'], row['name_short'],
                               row['name_full'], row['owned_by'], datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            elif table == 'operations':
                for _, row in df.iterrows():
                    c.execute('''INSERT INTO operations 
                                 (medicine_id, location_id, operation_type, operation_date, quantity, created_date) 
                                 VALUES (%s, %s, %s, %s, %s, %s)''',
                              (row['medicine_id'], row['location_id'], row['operation_type'], row['operation_date'],
                               row['quantity'], datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            conn.commit()
            log_action(f"Imported data into {table}", f"Rows: {len(df)}")
            st.success(f"Импортировано {len(df)} записей в таблицу {table}")
        else:
            st.error("Поддерживаются только CSV и Excel файлы")
    except Exception as e:
        st.error(f"Ошибка импорта: {e}")
    finally:
        conn.close()

def export_data(table):
    conn = get_db_connection()
    if conn is None:
        return
    df = pd.read_sql_query(f"SELECT * FROM {table}", conn)
    conn.close()
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name=table, index=False)
    log_action(f"Exported data from {table}", f"Rows: {len(df)}")
    st.download_button(label=f"Экспорт {table}", data=output.getvalue(),
                      file_name=f"{table}_export.xlsx",
                      mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# Валидация данных
def validate_medication_data(name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code):
    errors = []
    if not name:
        errors.append("Название не может быть пустым")
    if not gtin or len(gtin) > 20:
        errors.append("GTIN должен быть непустым и не длиннее 20 символов")
    if not sku or len(sku) > 20:
        errors.append("SKU должен быть непустым и не длиннее 20 символов")
    if not market:
        errors.append("Рынок не может быть пустым")
    if not batch_number:
        errors.append("Номер партии не может быть пустым")
    if not expiration_date:
        errors.append("Срок годности обязателен")
    if not dosage_form:
        errors.append("Форма выпуска обязательна")
    if not active_ingredient:
        errors.append("Активный ингредиент обязателен")
    if not package_size:
        errors.append("Объем/Размер упаковки обязателен")
    if not owned_by:
        errors.append("Компания-владелец обязательна")
    if atc_code and not re.match(r'^[A-Z]{1,2}[0-9]{2}[A-Z]{0,2}[0-9]{0,2}$', atc_code):
        errors.append("Код АТС имеет неверный формат (пример: A10BA02)")
    return errors

def validate_company_data(gln, name_short, name_full, gcp_compliant, registration_country, address, type):
    errors = []
    if not name_short:
        errors.append("Краткое название не может быть пустым")
    if len(name_short) > 50:
        errors.append("Краткое название не должно превышать 50 символов")
    if not name_full:
        errors.append("Полное название не может быть пустым")
    if len(name_full) > 100:
        errors.append("Полное название не должно превышать 100 символов")
    if gln and len(gln) > 20:
        errors.append("GLN не должен превышать 20 символов")
    if registration_country and len(registration_country) > 50:
        errors.append("Страна регистрации не должна превышать 50 символов")
    if address and len(address) > 200:
        errors.append("Адрес не должен превышать 200 символов")
    if type and len(type) > 50:
        errors.append("Тип не должен превышать 50 символов")
    return errors

def validate_location_data(gln, country, address, role, name_short, name_full, owned_by):
    errors = []
    if not address:
        errors.append("Адрес не может быть пустым")
    if len(address) > 200:
        errors.append("Адрес не должен превышать 200 символов")
    if gln and len(gln) > 20:
        errors.append("GLN не должен превышать 20 символов")
    if country and len(country) > 50:
        errors.append("Страна не должна превышать 50 символов")
    if role and len(role) > 50:
        errors.append("Роль не должна превышать 50 символов")
    if name_short and len(name_short) > 50:
        errors.append("Краткое название не должно превышать 50 символов")
    if name_full and len(name_full) > 100:
        errors.append("Полное название не должно превышать 100 символов")
    if not owned_by:
        errors.append("Компания-владелец обязательна")
    return errors

def validate_operation_data(medicine_id, location_id, operation_type, operation_date, quantity):
    errors = []
    if not medicine_id:
        errors.append("ID Препарата обязателен")
    if not location_id:
        errors.append("ID локации обязателен")
    if not operation_type:
        errors.append("Тип операции не может быть пустым")
    if not operation_date:
        errors.append("Дата операции обязательна")
    if not quantity or quantity <= 0:
        errors.append("Количество должно быть больше 0")
    return errors

# Функция авторизации
def login(username, password):
    conn = get_db_connection()
    if conn is None:
        return False
    c = conn.cursor()
    try:
        c.execute("SELECT password, role FROM users WHERE login = %s", (username,))
        result = c.fetchone()
        if result and result[0] == password:
            st.session_state['logged_in'] = True
            st.session_state['username'] = username
            st.session_state['role'] = result[1]
            log_action(f"Успешный вход пользователя: {username}")
            return True
        log_action(f"Неуспешная попытка входа: {username}")
        return False
    except psycopg2.Error as e:
        st.error(f"Ошибка базы данных: {e}")
        return False
    finally:
        conn.close()

# Интерфейс авторизации
def auth_interface():
    st.markdown("""
    <style>
        .login-container h2 {
            color: #333;
            margin-bottom: 20px;
            font-size: 28px;
        }
        .stButton > button {
            background-color: #4CAF50;
            color: white;
            padding: 12px 24px;
            font-size: 18px;
            border-radius: 5px;
            border: none;
            cursor: pointer;
            width: 100%;
        }
        .stButton > button:hover {
            background-color: #45a049;
        }
        .error-message {
            color: #d32f2f;
            font-size: 16px;
            margin-top: 10px;
        }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="login-container">', unsafe_allow_html=True)
    st.markdown('<h2>Авторизация в Kvinta</h2>', unsafe_allow_html=True)
    
    username = st.text_input("Логин", key="login_username", placeholder="Введите логин")
    password = st.text_input("Пароль", type="password", key="login_password", placeholder="Введите пароль")
    
    if st.button("Войти"):
        if login(username, password):
            st.success("Успешный вход!")
            st.session_state['show_kvinta_page'] = True
            st.session_state['show_main_page'] = False
            st.rerun()
        else:
            st.session_state['show_access_denied'] = True
            st.rerun()

# Страница отказа в доступе
def show_access_denied():
    st.markdown(""" 
    <style>
        .access-denied-container {
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            background-color: rgba(211, 47, 47, 0.5); /* Более прозрачный фон */
            padding: 20px;
            border-radius: 10px;
            text-align: center;
            color: white;
            font-size: 24px;
            max-width: 400px; /* Ограничение ширины */
            margin: 50px auto; /* Центрирование с отступом сверху */
        }
        .access-denied-text {
            font-size: 16px;
            margin-top: 10px;
        }
        .return-button {
            display: flex;
            justify-content: flex-start;
            margin-top: 20px;
            padding-left: 20px;
        }
        .return-button > button {
            background-color: #4CAF50;
            color: white;
            padding: 10px 20px;
            font-size: 16px;
            border-radius: 5px;
            border: none;
            cursor: pointer;
        }
        .return-button > button:hover {
            background-color: #45a049;
        }
    </style>
    """, unsafe_allow_html=True)
    st.markdown("""
    <div class="access-denied-container">
        Нет прав пользования подсистемой
        <div class="access-denied-text">Обратитесь к администратору для получения доступа.</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown('<div class="return-button">', unsafe_allow_html=True)
    if st.button("Вернуться к авторизации"):
        st.session_state['show_access_denied'] = False
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# Интерфейс и страницы
def show_view_data():
    st.subheader("Просмотр данных")
    entity = st.selectbox("Выберите тип данных", ["Препараты", "Компании", "Локации", "Операции"])
    
    if entity == "Препараты":
        df = get_medications()
        companies = get_companies()
        if not df.empty:
            if not companies.empty:
                # Merge с явным указанием суффиксов для избежания конфликтов
                df = df.merge(companies[['id', 'name_full']], left_on='owned_by', right_on='id', how='left', suffixes=('', '_company'))
                df = df.rename(columns={'name_full': 'owned_by_name'})
                # Удаляем только owned_by и id_company (если есть)
                display_df = df.drop(columns=['owned_by', 'id_company'], errors='ignore')
            else:
                # Если companies пустой, создаём owned_by_name с None
                df['owned_by_name'] = None
                display_df = df.drop(columns=['owned_by'], errors='ignore')
            # Проверяем наличие столбцов перед выбором
            available_columns = [col for col in ['id', 'owned_by_name', 'name', 'gtin', 'sku', 'atc_code', 'market', 'shared', 'batch_number', 'expiration_date', 'dosage_form', 'active_ingredient', 'package_size', 'created_date'] if col in display_df.columns]
            display_df = display_df[available_columns]
        else:
            display_df = df
    elif entity == "Компании":
        display_df = get_companies()
    elif entity == "Локации":
        df = get_locations()
        companies = get_companies()
        if not df.empty:
            if not companies.empty:
                df = df.merge(companies[['id', 'name_full']], left_on='owned_by', right_on='id', how='left', suffixes=('', '_company'))
                df = df.rename(columns={'name_full': 'owned_by_name'})
                display_df = df.drop(columns=['owned_by', 'id_company'], errors='ignore')
            else:
                df['owned_by_name'] = None
                display_df = df.drop(columns=['owned_by'], errors='ignore')
            available_columns = [col for col in ['id', 'owned_by_name', 'gln', 'country', 'address', 'role', 'name_short', 'name_full', 'created_date'] if col in display_df.columns]
            display_df = display_df[available_columns]
        else:
            display_df = df
    else:  # Операции
        df = get_operations()
        medicines = get_medications()
        locations = get_locations()
        if not df.empty:
            if not medicines.empty:
                df = df.merge(medicines[['id', 'name']], left_on='medicine_id', right_on='id', how='left', suffixes=('', '_med'))
                df = df.rename(columns={'name': 'medicine_name'})
            else:
                df['medicine_name'] = None
            if not locations.empty:
                df = df.merge(locations[['id', 'name_short']], left_on='location_id', right_on='id', how='left', suffixes=('', '_loc'))
                df = df.rename(columns={'name_short': 'location_name'})
            else:
                df['location_name'] = None
            display_df = df.drop(columns=['medicine_id', 'location_id', 'id_med', 'id_loc'], errors='ignore')
            available_columns = [col for col in ['id', 'medicine_name', 'location_name', 'operation_type', 'operation_date', 'quantity', 'created_date'] if col in display_df.columns]
            display_df = display_df[available_columns]
        else:
            display_df = df

    if not display_df.empty:
        st.write("### Данные")
        st.dataframe(display_df)
        st.write("### Числовая статистика")
        numeric_df = display_df.select_dtypes(include=['int64', 'float64'])
        if not numeric_df.empty:
            st.dataframe(numeric_df.describe())
        else:
            st.info("Нет числовых данных для статистики.")
        st.write("### Категориальная статистика")
        categorical_df = display_df.select_dtypes(include=['object', 'bool'])
        if not categorical_df.empty:
            st.dataframe(categorical_df.describe())
        else:
            st.info("Нет категориальных данных для статистики.")
        if st.button("Экспорт", key=f"export_{entity.lower()}_data"):
            export_data(entity.lower())
    else:
        st.warning(f"Нет данных для отображения. Добавьте {entity.lower()} на странице 'Добавить'.")

def show_edit_delete_data():
    if st.session_state['role'] not in ['admin', 'analyst']:
        log_action("Access denied to edit data", username=st.session_state['username'])
        st.error("Доступ запрещен")
        return
    st.subheader("Редактировать или удалить запись")
    action = st.radio("Выберите действие", ["Редактировать", "Удалить"], horizontal=True)
    entity = st.selectbox("Выберите тип записи", ["Препараты", "Компании", "Локации", "Операции"])

    if action == "Удалить":
        if entity == "Препараты":
            med_id = st.number_input("ID Препарата для удаления", min_value=1)
            if st.button("Удалить"):
                delete_medication(med_id)
                st.success("Препарат удален!")
        elif entity == "Компании":
            company_id = st.number_input("ID компании для удаления", min_value=1)
            if st.button("Удалить"):
                delete_company(company_id)
                st.success("Компания удалена!")
        elif entity == "Локации":
            location_id = st.number_input("ID локации для удаления", min_value=1)
            if st.button("Удалить"):
                delete_location(location_id)
                st.success("Локация удалена!")
        else:
            operation_id = st.number_input("ID операции для удаления", min_value=1)
            if st.button("Удалить"):
                delete_operation(operation_id)
                st.success("Операция удалена!")
    else:
        record_id = st.number_input("ID записи для редактирования", min_value=1)
        conn = get_db_connection()
        if conn is None:
            return
        c = conn.cursor()
        try:
            if entity == "Препараты":
                c.execute("SELECT * FROM medicines WHERE id = %s", (record_id,))
                record = c.fetchone()
                if record:
                    df = pd.DataFrame([record], columns=['id', 'owned_by', 'name', 'gtin', 'sku', 'market', 'shared', 'batch_number', 'expiration_date', 'dosage_form', 'active_ingredient', 'package_size', 'atc_code', 'created_date'])
                    companies = get_companies()
                    company_options = {f"{row['name_full']} (ID: {row['id']})": row['id'] for _, row in companies.iterrows()} if not companies.empty else {"Нет компаний": None}
                    with st.form(key=f"edit_med_{record_id}"):
                        name = st.text_input("Название", value=df['name'].iloc[0] or "")
                        gtin = st.text_input("GTIN", value=df['gtin'].iloc[0] or "")
                        sku = st.text_input("SKU", value=df['sku'].iloc[0] or "")
                        market = st.text_input("Рынок", value=df['market'].iloc[0] or "")
                        batch_number = st.text_input("Номер партии", value=df['batch_number'].iloc[0] or "")
                        expiration_date = st.date_input("Срок годности", value=pd.to_datetime(df['expiration_date'].iloc[0]) if pd.notnull(df['expiration_date'].iloc[0]) else None)
                        dosage_form = st.text_input("Форма выпуска", value=df['dosage_form'].iloc[0] or "")
                        active_ingredient = st.text_input("Активный ингредиент", value=df['active_ingredient'].iloc[0] or "")
                        package_size = st.text_input("Объем/Размер упаковки", value=df['package_size'].iloc[0] or "")
                        atc_code = st.text_input("Код АТС", value=df['atc_code'].iloc[0] or "" if pd.notnull(df['atc_code'].iloc[0]) else "")
                        owned_by_choice = st.selectbox("Компания-владелец", list(company_options.keys()), index=list(company_options.keys()).index(next((k for k, v in company_options.items() if v == df['owned_by'].iloc[0]), list(company_options.keys())[0])))
                        if st.form_submit_button("Сохранить"):
                            errors = validate_medication_data(name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, company_options[owned_by_choice], atc_code)
                            if errors:
                                for error in errors:
                                    st.error(error)
                            else:
                                c.execute("SELECT id FROM medicines WHERE gtin = %s AND sku = %s AND id != %s", (gtin, sku, record_id))
                                if c.fetchone():
                                    st.error("Препарат с таким GTIN и SKU уже существует")
                                else:
                                    edit_medication(record_id, name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, company_options[owned_by_choice], atc_code or None, st.session_state['username'])
                                    st.success("Препарат обновлен!")
                                    log_action("Edited medication", f"ID: {record_id}", st.session_state['username'])
                else:
                    st.error("Препарат с таким ID не найден")
            elif entity == "Компании":
                c.execute("SELECT * FROM companies WHERE id = %s", (record_id,))
                record = c.fetchone()
                if record:
                    df = pd.DataFrame([record], columns=['id', 'gln', 'name_short', 'name_full', 'gcp_compliant', 'registration_country', 'address', 'type'])
                    with st.form(key=f"edit_comp_{record_id}"):
                        gln = st.text_input("GLN", value=df['gln'].iloc[0] or "")
                        name_short = st.text_input("Краткое название", value=df['name_short'].iloc[0] or "")
                        name_full = st.text_input("Полное название", value=df['name_full'].iloc[0] or "")
                        gcp_compliant = st.checkbox("GCP-совместимость", value=df['gcp_compliant'].iloc[0] if pd.notnull(df['gcp_compliant'].iloc[0]) else False)
                        registration_country = st.text_input("Страна регистрации", value=df['registration_country'].iloc[0] or "")
                        address = st.text_input("Адрес", value=df['address'].iloc[0] or "")
                        type = st.text_input("Тип", value=df['type'].iloc[0] or "")
                        if st.form_submit_button("Сохранить"):
                            errors = validate_company_data(gln, name_short, name_full, gcp_compliant, registration_country, address, type)
                            if errors:
                                for error in errors:
                                    st.error(error)
                            else:
                                c.execute("SELECT id FROM companies WHERE gln = %s AND name_full = %s AND id != %s", (gln, name_full, record_id))
                                if c.fetchone():
                                    st.error("Компания с таким GLN и полным названием уже существует")
                                else:
                                    edit_company(record_id, gln or None, name_short, name_full, gcp_compliant, registration_country or None, address or None, type or None, st.session_state['username'])
                                    st.success("Компания обновлена!")
                                    log_action("Edited company", f"ID: {record_id}", st.session_state['username'])
                else:
                    st.error("Компания с таким ID не найдена")
            elif entity == "Локации":
                c.execute("SELECT * FROM locations WHERE id = %s", (record_id,))
                record = c.fetchone()
                if record:
                    df = pd.DataFrame([record], columns=['id', 'owned_by', 'gln', 'country', 'address', 'role', 'name_short', 'name_full', 'created_date'])
                    companies = get_companies()
                    company_options = {f"{row['name_full']} (ID: {row['id']})": row['id'] for _, row in companies.iterrows()} if not companies.empty else {"Нет компаний": None}
                    with st.form(key=f"edit_loc_{record_id}"):
                        gln = st.text_input("GLN", value=df['gln'].iloc[0] or "")
                        country = st.text_input("Страна", value=df['country'].iloc[0] or "")
                        address = st.text_input("Адрес", value=df['address'].iloc[0] or "")
                        role = st.text_input("Роль", value=df['role'].iloc[0] or "")
                        name_short = st.text_input("Краткое название", value=df['name_short'].iloc[0] or "")
                        name_full = st.text_input("Полное название", value=df['name_full'].iloc[0] or "")
                        owned_by_choice = st.selectbox("Компания-владелец", list(company_options.keys()), index=list(company_options.keys()).index(next((k for k, v in company_options.items() if v == df['owned_by'].iloc[0]), list(company_options.keys())[0])))
                        if st.form_submit_button("Сохранить"):
                            errors = validate_location_data(gln, country, address, role, name_short, name_full, company_options[owned_by_choice])
                            if errors:
                                for error in errors:
                                    st.error(error)
                            else:
                                edit_location(record_id, gln or None, country or None, address, role or None, name_short or None, name_full or None, company_options[owned_by_choice], st.session_state['username'])
                                st.success("Локация обновлена!")
                                log_action("Edited location", f"ID: {record_id}", st.session_state['username'])
                else:
                    st.error("Локация с таким ID не найдена")
            elif entity == "Операции":
                c.execute("SELECT * FROM operations WHERE id = %s", (record_id,))
                record = c.fetchone()
                if record:
                    df = pd.DataFrame([record], columns=['id', 'medicine_id', 'location_id', 'operation_type', 'operation_date', 'quantity', 'created_date'])
                    medicines = get_medications()
                    locations = get_locations()
                    medicine_options = {f"{row['name']} (ID: {row['id']})": row['id'] for _, row in medicines.iterrows()} if not medicines.empty else {"Нет препаратов": None}
                    location_options = {f"{row['name_short'] or row['name_full'] or f'Локация ID {row['id']}'} (ID: {row['id']})": row['id'] for _, row in locations.iterrows()} if not locations.empty else {"Нет локаций": None}
                    with st.form(key=f"edit_op_{record_id}"):
                        medicine_choice = st.selectbox("Препарат", list(medicine_options.keys()), index=list(medicine_options.keys()).index(next((k for k, v in medicine_options.items() if v == df['medicine_id'].iloc[0]), list(medicine_options.keys())[0])))
                        location_choice = st.selectbox("Локация", list(location_options.keys()), index=list(location_options.keys()).index(next((k for k, v in location_options.items() if v == df['location_id'].iloc[0]), list(location_options.keys())[0])))
                        operation_type = st.selectbox("Тип операции", ["Агрегация", "Дистрибьютор", "Поставка", "Списание", "Производство", "Перемещение"], index=["Агрегация", "Дистрибьютор", "Поставка", "Списание", "Производство", "Перемещение"].index(df['operation_type'].iloc[0]) if df['operation_type'].iloc[0] in ["Агрегация", "Дистрибьютор", "Поставка", "Списание", "Производство", "Перемещение"] else 0)
                        operation_date = st.date_input("Дата операции", value=pd.to_datetime(df['operation_date'].iloc[0]) if pd.notnull(df['operation_date'].iloc[0]) else None)
                        quantity = st.number_input("Количество", min_value=1, value=int(df['quantity'].iloc[0]) if pd.notnull(df['quantity'].iloc[0]) else 1)
                        if st.form_submit_button("Сохранить"):
                            errors = validate_operation_data(medicine_options[medicine_choice], location_options[location_choice], operation_type, operation_date, quantity)
                            if errors:
                                for error in errors:
                                    st.error(error)
                            else:
                                edit_operation(record_id, medicine_options[medicine_choice], location_options[location_choice], operation_type, operation_date, quantity, st.session_state['username'])
                                st.success("Операция обновлена!")
                                log_action("Edited operation", f"ID: {record_id}", st.session_state['username'])
                else:
                    st.error("Операция с таким ID не найдена")
        except Exception as e:
            st.error(f"Ошибка редактирования: {e}")
            log_action("Edit error", f"Entity: {entity}, ID: {record_id}, Error: {str(e)}", st.session_state['username'])
        finally:
            conn.close()

def show_add_data():
    st.subheader("Добавить новую запись")
    entity = st.selectbox("Выберите тип записи", ["Препараты", "Компании", "Локации", "Операции"])
    if entity == "Препараты":
        companies = get_companies()
        company_options = {f"{row['name_full']} (ID: {row['id']})": row['id'] for _, row in companies.iterrows()} if not companies.empty else {"Нет компаний": None}
        with st.form(key="add_med"):
            name = st.text_input("Название")
            gtin = st.text_input("GTIN")
            sku = st.text_input("SKU")
            market = st.text_input("Рынок")
            batch_number = st.text_input("Номер партии")
            expiration_date = st.date_input("Срок годности")
            dosage_form = st.text_input("Форма выпуска")
            active_ingredient = st.text_input("Активный ингредиент")
            package_size = st.text_input("Объем/Размер упаковки")
            atc_code = st.text_input("Код АТС (например, A10BA02)")
            owned_by_choice = st.selectbox("Компания-владелец", list(company_options.keys()))
            uploaded_file = st.file_uploader("Импорт из CSV/Excel", type=['csv', 'xlsx'], key="med_import")
            if uploaded_file:
                import_data(uploaded_file)
            if st.form_submit_button("Добавить"):
                owned_by = company_options.get(owned_by_choice)
                errors = validate_medication_data(name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code)
                if errors:
                    for error in errors:
                        st.error(error)
                else:
                    conn = get_db_connection()
                    if conn:
                        c = conn.cursor()
                        c.execute('''SELECT id FROM medicines 
                                    WHERE name = %s AND gtin = %s AND sku = %s AND market = %s 
                                    AND batch_number = %s AND expiration_date = %s AND dosage_form = %s 
                                    AND active_ingredient = %s AND package_size = %s AND owned_by = %s AND atc_code = %s''',
                                (name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code))
                        if c.fetchone():
                            st.error("Такая запись уже существует")
                            conn.close()
                        else:
                            conn.close()
                            add_medication(name, gtin, sku, market, batch_number, expiration_date, dosage_form, active_ingredient, package_size, owned_by, atc_code, st.session_state['username'])
                            st.success("Препарат добавлен!")
    elif entity == "Компании":
        with st.form(key="add_comp"):
            gln = st.text_input("GLN")
            name_short = st.text_input("Краткое название")
            name_full = st.text_input("Полное название")
            gcp_compliant = st.checkbox("GCP-совместимость")
            registration_country = st.text_input("Страна регистрации")
            address = st.text_input("Адрес")
            type = st.text_input("Тип")
            uploaded_file = st.file_uploader("Импорт из CSV/Excel", type=['csv', 'xlsx'], key="comp_import")
            if uploaded_file:
                import_data(uploaded_file)
            if st.form_submit_button("Добавить"):
                errors = validate_company_data(gln, name_short, name_full, gcp_compliant, registration_country, address, type)
                if errors:
                    for error in errors:
                        st.error(error)
                else:
                    conn = get_db_connection()
                    if conn:
                        c = conn.cursor()
                        c.execute('''SELECT id FROM companies 
                                     WHERE gln = %s AND name_short = %s AND name_full = %s 
                                     AND gcp_compliant = %s AND registration_country = %s 
                                     AND address = %s AND type = %s''',
                                  (gln, name_short, name_full, gcp_compliant, registration_country, address, type))
                        if c.fetchone():
                            st.error("Такая запись уже существует")
                            conn.close()
                        else:
                            conn.close()
                            add_company(gln, name_short, name_full, gcp_compliant, registration_country, address, type, st.session_state['username'])
                            st.success("Компания добавлена!")
    elif entity == "Локации":
        companies = get_companies()
        company_options = {f"{row['name_full']} (ID: {row['id']})": row['id'] for _, row in companies.iterrows()} if not companies.empty else {"Нет компаний": None}
        with st.form(key="add_loc"):
            gln = st.text_input("GLN")
            country = st.text_input("Страна")
            address = st.text_input("Адрес")
            role = st.text_input("Роль")
            name_short = st.text_input("Краткое название")
            name_full = st.text_input("Полное название")
            owned_by_choice = st.selectbox("Компания-владелец", list(company_options.keys()))
            uploaded_file = st.file_uploader("Импорт из CSV/Excel", type=['csv', 'xlsx'], key="loc_import")
            if uploaded_file:
                import_data(uploaded_file)
            if st.form_submit_button("Добавить"):
                owned_by = company_options.get(owned_by_choice)
                errors = validate_location_data(gln, country, address, role, name_short, name_full, owned_by)
                if errors:
                    for error in errors:
                        st.error(error)
                else:
                    add_location(gln, country, address, role, name_short, name_full, owned_by, st.session_state['username'])
                    st.success("Локация добавлена!")
    else:
        medicines = get_medications()
        locations = get_locations()
        medicine_options = {f"{row['name']} (ID: {row['id']})": row['id'] for _, row in medicines.iterrows()} if not medicines.empty else {"Нет препаратов": None}
        location_options = {f"{row['name_short'] or row['name_full'] or f'Локация ID {row['id']}'} (ID: {row['id']})": row['id'] for _, row in locations.iterrows()} if not locations.empty else {"Нет локаций": None}
        with st.form(key="add_op"):
            medicine_choice = st.selectbox("Препарат", list(medicine_options.keys()))
            location_choice = st.selectbox("Локация", list(location_options.keys()))
            operation_type = st.selectbox("Тип операции", ["Агрегация", "Дистрибьютор", "Поставка", "Списание", "Производство", "Перемещение"])
            operation_date = st.date_input("Дата операции")
            quantity = st.number_input("Количество", min_value=1, value=1)
            uploaded_file = st.file_uploader("Импорт из CSV/Excel", type=['csv', 'xlsx'], key="op_import")
            if uploaded_file:
                import_data(uploaded_file)
            if st.form_submit_button("Добавить"):
                medicine_id = medicine_options.get(medicine_choice)
                location_id = location_options.get(location_choice)
                errors = validate_operation_data(medicine_id, location_id, operation_type, operation_date, quantity)
                if errors:
                    for error in errors:
                        st.error(error)
                else:
                    add_operation(medicine_id, location_id, operation_type, operation_date, quantity, st.session_state['username'])
                    st.success("Операция добавлена!")

def show_filter_data():
    st.subheader("Фильтрация")
    entity = st.selectbox("Выберите тип данных", ["Препараты", "Компании", "Локации", "Операции"])
    if entity == "Препараты":
        df = get_medications()
        filter_options = {
            "name": "Название",
            "gtin": "GTIN",
            "sku": "SKU",
            "market": "Рынок",
            "batch_number": "Номер партии",
            "expiration_date": "Срок годности",
            "dosage_form": "Форма выпуска",
            "active_ingredient": "Активный ингредиент",
            "package_size": "Объем/Размер упаковки",
            "atc_code": "Код АТС",
            "created_date": "Дата создания",
            "owned_by": "ID компании-владельца"
        }
    elif entity == "Компании":
        df = get_companies()
        filter_options = {
            "gln": "GLN",
            "name_short": "Краткое название",
            "name_full": "Полное название",
            "gcp_compliant": "GCP-совместимость",
            "registration_country": "Страна регистрации",
            "address": "Адрес",
            "type": "Тип"
        }
    elif entity == "Локации":
        df = get_locations()
        filter_options = {
            "gln": "GLN",
            "country": "Страна",
            "address": "Адрес",
            "role": "Роль",
            "name_short": "Краткое название",
            "name_full": "Полное название",
            "owned_by": "ID компании-владельца",
            "created_date": "Дата создания"
        }
    else:
        df = get_operations()
        filter_options = {
            "medicine_id": "ID Препарата",
            "location_id": "ID локации",
            "operation_type": "Тип операции",
            "operation_date": "Дата операции",
            "quantity": "Количество",
            "created_date": "Дата создания"
        }
    if df.empty:
        st.warning(f"Нет данных для фильтрации. Добавьте {entity.lower()} на странице 'Добавить'.")
        return
    selected_filters = st.multiselect(
        "Выберите параметры для фильтрации",
        options=list(filter_options.keys()),
        format_func=lambda x: filter_options[x]
    )
    if not selected_filters:
        st.info("Выберите хотя бы один параметр для фильтрации.")
        return
    filter_values = {}
    for param in selected_filters:
        if param in ["name", "gtin", "sku", "market", "batch_number", "address", "gln", "country", "role", "name_short", "name_full", "registration_country", "type", "operation_type", "dosage_form", "active_ingredient", "package_size"]:
            filter_values[param] = st.text_input(f"Введите {filter_options[param]} (или оставьте пустым)")
        elif param in ["expiration_date", "operation_date", "created_date"]:
            filter_values[param] = st.date_input(f"Выберите {filter_options[param]}")
        elif param in ["quantity", "medicine_id", "location_id", "owned_by"]:
            filter_values[param] = st.number_input(f"Введите {filter_options[param]}", step=1)
        elif param == "gcp_compliant":
            filter_values[param] = st.selectbox(f"Выберите {filter_options[param]}", ["Любое", "Да", "Нет"])

    if st.button("Применить фильтр"):
        filtered_df = df.copy()
        for param in selected_filters:
            if param in ["name", "gtin", "sku", "market", "batch_number", "address", "gln", "country", "role", "name_short", "name_full", "registration_country", "type", "operation_type", "dosage_form", "active_ingredient", "package_size"]:
                if filter_values[param]:
                    filtered_df = filtered_df[filtered_df[param].str.contains(filter_values[param], case=False, na=False)]
            elif param in ["expiration_date", "operation_date", "created_date"]:
                filtered_df = filtered_df[filtered_df[param].astype(str).str.startswith(str(filter_values[param]), na=False)]
            elif param in ["quantity", "medicine_id", "location_id", "owned_by"]:
                filtered_df = filtered_df[filtered_df[param] == filter_values[param]]
            elif param == "gcp_compliant":
                if filter_values[param] != "Любое":
                    filtered_df = filtered_df[filtered_df[param] == (filter_values[param] == "Да")]
        if filtered_df.empty:
            st.warning("Нет данных, соответствующих выбранным фильтрам.")
        else:
            st.subheader("Отфильтрованные данные")
            st.dataframe(filtered_df)

def show_visualize():
    st.subheader("Визуализация данных")
    entity = st.selectbox("Выберите тип данных", ["Препараты", "Компании", "Локации", "Операции"])
    if entity == "Препараты":
        df = get_medications()
        if df.empty:
            st.warning("Нет данных для визуализации. Добавьте Препараты на странице 'Добавить'.")
            return
        viz_type = st.selectbox("Тип визуализации", [
            "Распределение по рынкам",
            "Доля препаратов по сроку годности",
            "Распределение по формам выпуска",
            "Препараты по размеру упаковки"
        ])
        if viz_type == "Распределение по рынкам":
            if 'market' in df.columns:
                fig = px.histogram(df, x='market', title="Распределение препаратов по рынкам", color='market')
                fig.update_layout(xaxis_title="Рынок", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'market' отсутствует в данных.")
        elif viz_type == "Доля препаратов по сроку годности":
            if 'expiration_date' in df.columns:
                current_date = pd.to_datetime("2025-05-21")
                df['days_to_expiry'] = (pd.to_datetime(df['expiration_date']) - current_date).dt.days
                df['expiry_status'] = pd.cut(df['days_to_expiry'],
                                            bins=[-float('inf'), 0, 180, 365, float('inf')],
                                            labels=['Просрочено', 'Менее 6 мес.', '6-12 мес.', 'Более года'])
                fig = px.pie(df, names='expiry_status', title="Доля препаратов по сроку годности")
                fig.update_layout(legend_title="Статус срока годности", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'expiration_date' отсутствует в данных.")
        elif viz_type == "Распределение по формам выпуска":
            if 'dosage_form' in df.columns:
                fig = px.histogram(df, x='dosage_form', title="Распределение препаратов по формам выпуска", color='dosage_form')
                fig.update_layout(xaxis_title="Форма выпуска", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'dosage_form' отсутствует в данных.")
        else:
            if 'package_size' in df.columns:
                fig = px.histogram(df, x='package_size', title="Препараты по размеру упаковки", color='package_size')
                fig.update_layout(xaxis_title="Объем/Размер упаковки", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'package_size' отсутствует в данных.")
    elif entity == "Компании":
        df = get_companies()
        if df.empty:
            st.warning("Нет данных для визуализации. Добавьте компании на странице 'Добавить'.")
            return
        viz_type = st.selectbox("Тип визуализации", [
            "Распределение по странам регистрации",
            "Доля по типам компаний",
            "Компании по GCP-совместимости"
        ])
        if viz_type == "Распределение по странам регистрации":
            if 'registration_country' in df.columns:
                fig = px.histogram(df, x='registration_country', title="Распределение компаний по странам регистрации", color='registration_country')
                fig.update_layout(xaxis_title="Страна регистрации", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'registration_country' отсутствует в данных.")
        elif viz_type == "Доля по типам компаний":
            if 'type' in df.columns:
                fig = px.pie(df, names='type', title="Доля компаний по типам")
                fig.update_layout(legend_title="Тип компании", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'type' отсутствует в данных.")
        else:
            if 'gcp_compliant' in df.columns:
                fig = px.histogram(df, x='gcp_compliant', title="Компании по GCP-совместимости", color='gcp_compliant')
                fig.update_layout(xaxis_title="GCP-совместимость", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'gcp_compliant' отсутствует в данных.")
    elif entity == "Локации":
        df = get_locations()
        if df.empty:
            st.warning("Нет данных для визуализации. Добавьте локации на странице 'Добавить'.")
            return
        viz_type = st.selectbox("Тип визуализации", [
            "Распределение по странам",
            "Распределение по ролям",
            "Локации по компаниям"
        ])
        if viz_type == "Распределение по странам":
            if 'country' in df.columns:
                fig = px.histogram(df, x='country', title="Распределение локаций по странам", color='country')
                fig.update_layout(xaxis_title="Страна", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'country' отсутствует в данных.")
        elif viz_type == "Распределение по ролям":
            if 'role' in df.columns:
                fig = px.histogram(df, x='role', title="Распределение локаций по ролям", color='role')
                fig.update_layout(xaxis_title="Роль", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'role' отсутствует в данных.")
        else:
            if 'owned_by' in df.columns:
                fig = px.histogram(df, x='owned_by', title="Локации по компаниям", color='owned_by')
                fig.update_layout(xaxis_title="ID компании", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'owned_by' отсутствует в данных.")
    else:
        df = get_operations()
        if df.empty:
            st.warning("Нет данных для визуализации. Добавьте операции на странице 'Добавить'.")
            return
        viz_type = st.selectbox("Тип визуализации", [
            "Количество операций по датам",
            "Доля по типам операций",
            "Количество по типам операций",
            "Операции по Препаратам"
        ])
        if viz_type == "Количество операций по датам":
            if 'operation_date' in df.columns and 'quantity' in df.columns:
                df['operation_date'] = pd.to_datetime(df['operation_date'])
                df_grouped = df.groupby(df['operation_date'].dt.date)['quantity'].sum().reset_index()
                fig = px.line(df_grouped, x='operation_date', y='quantity', title="Количество операций по датам")
                fig.update_layout(xaxis_title="Дата операции", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонки 'operation_date' или 'quantity' отсутствуют в данных.")
        elif viz_type == "Доля по типам операций":
            if 'operation_type' in df.columns:
                fig = px.pie(df, names='operation_type', title="Доля операций по типам")
                fig.update_layout(legend_title="Тип операции", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'operation_type' отсутствует в данных.")
        elif viz_type == "Количество по типам операций":
            if 'operation_type' in df.columns:
                fig = px.histogram(df, x='operation_type', title="Количество по типам операций", color='operation_type')
                fig.update_layout(xaxis_title="Тип операции", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'operation_type' отсутствует в данных.")
        else:
            if 'medicine_id' in df.columns:
                # Объединяем с таблицей medicines, чтобы получить названия препаратов
                medicines = get_medications()
                df = df.merge(medicines[['id', 'name']], left_on='medicine_id', right_on='id', how='left')
                df['medicine_name'] = df['name'].fillna('Не указан')
                fig = px.histogram(df, x='medicine_name', title="Операции по Препаратам", color='medicine_name')
                fig.update_layout(xaxis_title="Название препарата", yaxis_title="Количество", showlegend=True, legend=dict(orientation="v", yanchor="top", y=1, xanchor="right", x=1))
                st.plotly_chart(fig)
            else:
                st.error("Колонка 'medicine_id' отсутствует в данных.")

def show_reports():
    if st.session_state['role'] not in ['admin', 'analyst']:
        st.error("Доступ запрещен")
        return
    st.subheader("Создание отчетов")

    medicines = get_medications()
    companies = get_companies()
    locations = get_locations()
    operations = get_operations()

    medicine_options = {f"{row['name']} (ID: {row['id']})": row['id'] for _, row in medicines.iterrows()} if not medicines.empty else {"Нет препаратов": None}
    with st.form(key="report_form"):
        report_title = st.text_input("Название отчета", placeholder="Введите название отчета")
        medicine_choice = st.selectbox("Препарат", list(medicine_options.keys()), help="Выберите препарат для отчета")
        submit_button = st.form_submit_button("Сформировать отчет")

    if submit_button:
        if not report_title:
            st.error("Название отчета обязательно")
            return
        if not medicine_choice or not medicine_options[medicine_choice]:
            st.error("Выберите препарат")
            return

        # Фильтрация данных
        med_id = medicine_options[medicine_choice]
        filtered_meds = medicines[medicines['id'] == med_id]
        filtered_ops = operations[operations['medicine_id'] == med_id]
        company_id = filtered_meds['owned_by'].iloc[0] if not filtered_meds.empty else None
        filtered_companies = companies[companies['id'] == company_id] if company_id else pd.DataFrame()
        location_ids = filtered_ops['location_id'].unique() if not filtered_ops.empty else []
        filtered_locations = locations[locations['id'].isin(location_ids)] if location_ids else pd.DataFrame()

        # Создание отчета
        doc = Document()
        # Настройка стилей для компактности
        styles = doc.styles
        compact_style = styles.add_style('Compact', WD_STYLE_TYPE.PARAGRAPH)
        compact_style.font.size = Pt(9)  # Меньший шрифт
        compact_style.paragraph_format.space_after = Pt(2)  # Минимальный отступ после абзаца
        compact_style.paragraph_format.line_spacing = 1.0  # Одинарный межстрочный интервал

        doc.add_heading("KVINTA (отчеты)", 0)
        doc.add_heading(f"Отчет: {report_title}", 1)

        # Препарат
        doc.add_heading("Препарат", level=2)
        if not filtered_meds.empty:
            row = filtered_meds.iloc[0]
            fields = [
                ("Название", str(row.get('name', 'Не указано'))),
                ("GTIN", str(row.get('gtin', 'Не указано'))),
                ("SKU", str(row.get('sku', 'Не указано'))),
                ("Рынок", str(row.get('market', 'Не указано'))),
                ("Партия", str(row.get('batch_number', 'Не указано'))),
                ("Срок годности", str(row.get('expiration_date', 'Не указано'))),
                ("Форма", str(row.get('dosage_form', 'Не указано'))),
                ("Ингредиент", str(row.get('active_ingredient', 'Не указано'))),
                ("Упаковка", str(row.get('package_size', 'Не указано'))),
                ("Код АТС", str(row.get('atc_code', 'Не указано')))
            ]
            table = doc.add_table(rows=len(fields) + 1, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].style = compact_style
            # Заголовки
            headers = ["Параметр", "Значение"]
            table.rows[0].cells[0].text = headers[0]
            table.rows[0].cells[1].text = headers[1]
            # Данные
            for i, (header, value) in enumerate(fields, 1):
                cells = table.rows[i].cells
                cells[0].text = header
                cells[1].text = value
        else:
            p = doc.add_paragraph("Данные о препарате не найдены", style='Compact')

        # Компания
        doc.add_heading("Компания", level=2)
        if not filtered_companies.empty:
            row = filtered_companies.iloc[0]
            fields = [
                ("GLN", str(row.get('gln', 'Не указано'))),
                ("Краткое название", str(row.get('name_short', 'Не указано'))),
                ("Полное название", str(row.get('name_full', 'Не указано'))),
                ("GCP", "Да" if row.get('gcp_compliant', False) else "Нет"),
                ("Страна", str(row.get('registration_country', 'Не указано'))),
                ("Адрес", str(row.get('address', 'Не указано'))),
                ("Тип", str(row.get('type', 'Не указано')))
            ]
            table = doc.add_table(rows=len(fields) + 1, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].style = compact_style
            headers = ["Параметр", "Значение"]
            table.rows[0].cells[0].text = headers[0]
            table.rows[0].cells[1].text = headers[1]
            for i, (header, value) in enumerate(fields, 1):
                cells = table.rows[i].cells
                cells[0].text = header
                cells[1].text = value
        else:
            p = doc.add_paragraph("Компания не найдена", style='Compact')

        # Местоположение
        doc.add_heading("Местоположение", level=2)
        if not filtered_locations.empty:
            table = doc.add_table(rows=len(filtered_locations) + 1, cols=6)
            table.style = 'Table Grid'
            table.autofit = True
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].style = compact_style
            headers = ["GLN", "Страна", "Адрес", "Роль", "Краткое название", "Полное название"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            for idx, row in enumerate(filtered_locations.itertuples(), 1):
                cells = table.rows[idx].cells
                cells[0].text = str(row.gln)
                cells[1].text = str(row.country)
                cells[2].text = str(row.address)
                cells[3].text = str(row.role)
                cells[4].text = str(row.name_short)
                cells[5].text = str(row.name_full)
        else:
            p = doc.add_paragraph("Локации не найдены", style='Compact')

        # Операции
        doc.add_heading("Операции", level=2)
        if not filtered_ops.empty:
            table = doc.add_table(rows=len(filtered_ops) + 1, cols=5)
            table.style = 'Table Grid'
            table.autofit = True
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].style = compact_style
            headers = ["Препарат", "Локация", "Тип операции", "Дата", "Кол-во"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            for idx, row in enumerate(filtered_ops.itertuples(), 1):
                med_name = medicines[medicines['id'] == row.medicine_id]['name'].iloc[0] if row.medicine_id in medicines['id'].values else 'Не указан'
                loc_name = locations[locations['id'] == row.location_id]['name_short'].iloc[0] if row.location_id in locations['id'].values else 'Не указана'
                cells = table.rows[idx].cells
                cells[0].text = med_name
                cells[1].text = loc_name
                cells[2].text = str(row.operation_type)
                cells[3].text = str(row.operation_date)
                cells[4].text = str(row.quantity)
        else:
            p = doc.add_paragraph("Операции не найдены", style='Compact')
        # Сохранение в Word
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        st.download_button(
            label="Скачать отчет (Word)",
            data=word_buffer,
            file_name=f"{report_title}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        log_action("Generated report", f"Title: {report_title}", st.session_state['username'])
       
def show_logs():
    st.subheader("Просмотр логов (Админ)")
    if st.session_state['role'] != 'admin':
        log_action("Access denied to view logs", username=st.session_state['username'])
        st.error("Доступ запрещен")
        return
    st.write("### Основные логи")
    logs = get_logs('main')
    st.text_area("Логи", value="".join(logs), height=400)
    st.download_button(
        label="Скачать основные логи",
        data="".join(logs),
        file_name="pharma_metadata.log",
        mime="text/plain"
    )
    # st.write("### Логи отказов в доступе к редактированию")
    # edit_logs = get_logs('edit')
    # st.text_area("Логи отказов", value="".join(edit_logs), height=200)
    # st.download_button(
        #label="Скачать логи отказов",
        #data="".join(edit_logs),
        #file_name="edit_access_denied.log",
        #mime="text/plain"
    #)

def show_kvinta_page():
    st.markdown("""
    <style>
        .kvinta-container {
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            padding: 30px;
            border-radius: 10px;
            text-align: center;
            background-color: #f0f2f6;
            max-width: 800px;
            margin: 50px auto;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        }
        .kvinta-container h1 {
            color: #2c3e50;
            font-size: 36px;
            margin-bottom: 20px;
            font-weight: bold;
        }
        .kvinta-container p {
            font-size: 18px;
            color: #555;
            margin-bottom: 20px;
            line-height: 1.6;
        }
        .button-container {
            display: flex;
            gap: 15px;
            flex-wrap: wrap;
            justify-content: center;
        }
        .stButton > button {
            background-color: #4CAF50;
            color: white;
            padding: 12px 24px;
            font-size: 18px;
            border-radius: 5px;
            border: none;
            cursor: pointer;
            transition: background-color 0.3s;
        }
        .stButton > button:hover {
            background-color: #45a049;
        }
        .secondary-button > button {
            background-color: #6c757d;
            color: white;
            padding: 10px 20px;
            font-size: 16px;
            border-radius: 5px;
            border: none;
            cursor: pointer;
        }
        .secondary-button > button:hover {
            background-color: #5a6268;
        }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="kvinta-container">', unsafe_allow_html=True)
    st.markdown('<h1>Kvinta</h1>', unsafe_allow_html=True)
    st.markdown("""
    <p>Добро пожаловать в платформу Kvinta — ваш надежный партнер в управлении данными фармацевтической отрасли. 
    Наша система предоставляет передовые решения для управления метаданными, аналитики и логистики лекарственных препаратов. 
    Перейдите в подсистему управления метаданными для работы с данными или ознакомьтесь с дополнительными возможностями Kvinta.</p>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="button-container">', unsafe_allow_html=True)
    if st.button("Перейти к подсистеме управления метаданными", key="main_system_button"):
        st.session_state['show_kvinta_page'] = False
        st.session_state['show_main_page'] = True
        st.rerun()
    
    # Placeholder buttons for a more interactive feel (non-functional)
    st.markdown('<div class="secondary-button">', unsafe_allow_html=True)
    st.button("О платформе Kvinta", key="about_button")
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="secondary-button">', unsafe_allow_html=True)
    st.button("Наши услуги", key="services_button")
    st.markdown('</div>', unsafe_allow_html=True)
       
    st.markdown('<div class="secondary-button">', unsafe_allow_html=True)
    st.button("Связаться с нами", key="contact_button")
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    
def show_home():
    st.title("Pharma Metadata System")
    st.subheader("Главная страница")
    st.write("""
    ### Добро пожаловать в Pharma Metadata System
    - **Версия**: 0.01
    - **Разработчик**: Kvinta
    - **Описание**: Подсистема предназначена для работы с метаданными лекарственных препаратов.
    - **Функции**:
      - Просмотр и добавление записей.
      - Редактирование и удаление данных (для администраторов и аналитиков).
      - Фильтрация и визуализация данных.
      - Импорт и экспорт данных в форматах CSV.
      - Создание отчетов (для администраторов и аналитиков).
      - Логирование действий пользователей (для администраторов).
    """)
    if st.button("Выйти"):
        log_action("User exited subsystem", st.session_state['username'])
        st.session_state['show_kvinta_page'] = True
        st.session_state['show_main_page'] = False
        st.session_state['show_access_denied'] = False
        st.rerun()

def main():
    init_db()
    clear_logs_daily()

    st.markdown("""
    <style>
        body {
            font-size: 24px;
        }
        h1 {
            font-size: 25px;
        }
        h2 {
            font-size: 25px;
        }
        .stTextInput > div > div > input {
            font-size: 20px;
        }
        .stSelectbox > div > div > select {
            font-size: 20px;
        }
        .stDataFrame {
            font-size: 25px;
        }
        .stButton > button {
            font-size: 20px;
        }
        .footer {
            position: fixed;
            bottom: 15px;
            right: 15px;
            color: gray;
            font-size: 14px;
        }
    </style>
    """, unsafe_allow_html=True)

    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False
        st.session_state['role'] = None
        st.session_state['username'] = None
        st.session_state['show_access_denied'] = False
        st.session_state['show_kvinta_page'] = False
        st.session_state['show_main_page'] = False

    st.subheader("Pharma Metadata System")
    if st.session_state['show_access_denied']:
        show_access_denied()
    elif not st.session_state['logged_in']:
        auth_interface()
    elif st.session_state['show_kvinta_page']:
        show_kvinta_page()
    elif st.session_state['show_main_page']:
        if st.session_state['role'] in ['admin', 'operator', 'analyst']:
            if st.session_state['role'] == 'admin':
                menu = ["Главная страница", "Просмотр", "Добавить", "Редактировать", "Фильтрация", "Визуализация", "Отчеты", "Логи"]
            elif st.session_state['role'] == 'analyst':
                menu = ["Главная страница", "Просмотр", "Добавить", "Редактировать", "Фильтрация", "Визуализация", "Отчеты"]
            else:  # operator
                menu = ["Главная страница", "Просмотр", "Добавить"]

            st.sidebar.title(f"Добро пожаловать, {st.session_state['role']}")
            choice = st.sidebar.selectbox("Меню", menu, index=0)

            if choice == "Главная страница":
                show_home()
            elif choice == "Просмотр":
                show_view_data()
            elif choice == "Добавить":
                show_add_data()
            elif choice == "Редактировать":
                show_edit_delete_data()
            elif choice == "Фильтрация":
                show_filter_data()
            elif choice == "Визуализация":
                show_visualize()
            elif choice == "Отчеты":
                show_reports()
            elif choice == "Логи":
                show_logs()
        else:
            show_access_denied()
    else:
        # Default to Kvinta page after login if no other page is active
        show_kvinta_page()
            
            
if __name__ == "__main__":
    main()
